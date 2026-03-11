from __future__ import annotations

from pathlib import Path

from engine.rag.constants import SourceFormat
from engine.rag.ingest.loaders.base import BaseLoader, LoadedDocument, LoadedSection
from engine.shared.exceptions import RAGLoaderError
from engine.shared.logging import get_logger

logger = get_logger(__name__)


class DocxLoader(BaseLoader):
    @property
    def supported_format(self) -> SourceFormat:
        return SourceFormat.DOCX

    async def load(self, path: Path) -> LoadedDocument:
        if not self.can_load(path):
            raise RAGLoaderError(
                f"Cannot load docx file: {path}",
                details={"path": str(path)},
            )
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RAGLoaderError(
                "python-docx is required for DOCX loading",
                details={"path": str(path)},
            ) from exc

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            raise RAGLoaderError(
                f"Failed to parse docx file: {path}",
                details={"path": str(path), "error": str(exc)},
            ) from exc

        paragraphs: list[str] = []
        sections: list[LoadedSection] = []
        current_heading: str | None = None
        current_level: int = 0
        current_lines: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower()

            if style_name.startswith("heading"):
                if current_heading is not None:
                    sections.append(LoadedSection(
                        heading=current_heading,
                        level=current_level,
                        content="\n".join(current_lines).strip(),
                    ))
                    current_lines = []

                level_str = style_name.replace("heading", "").strip()
                current_level = int(level_str) if level_str.isdigit() else 1
                current_heading = text
            else:
                current_lines.append(text)
                paragraphs.append(text)

        if current_heading is not None:
            sections.append(LoadedSection(
                heading=current_heading,
                level=current_level,
                content="\n".join(current_lines).strip(),
            ))

        full_content = "\n\n".join(paragraphs)
        if not full_content.strip():
            raise RAGLoaderError(
                f"DOCX file has no extractable text: {path}",
                details={"path": str(path)},
            )

        title = path.stem.replace("_", " ").title()
        if sections and sections[0].level == 1:
            title = sections[0].heading

        logger.info(
            "loaded_docx",
            path=str(path),
            title=title,
            section_count=len(sections),
            paragraph_count=len(paragraphs),
        )

        return LoadedDocument(
            content=full_content,
            source_path=str(path),
            source_format=SourceFormat.DOCX,
            title=title,
            sections=tuple(sections),
        )
